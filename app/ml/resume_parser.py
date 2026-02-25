"""
Resume Parser: Extracts plain text from PDF, DOCX, and TXT files.
Uses PyMuPDF for PDF (fast, handles scanned PDFs) and python-docx for Word.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from uploaded resume file.
    Returns cleaned plain text string.
    """
    ext = Path(filename).suffix.lower()

    try:
        if ext == '.pdf':
            return _parse_pdf(file_bytes)
        elif ext == '.docx':
            return _parse_docx(file_bytes)
        elif ext == '.txt':
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logger.error(f"Resume parse error for {filename}: {e}")
        raise


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        text = '\n'.join(text_parts)
        return _clean_text(text)
    except ImportError:
        logger.warning("PyMuPDF not available, trying fallback PDF parser")
        return _parse_pdf_fallback(file_bytes)


def _parse_pdf_fallback(file_bytes: bytes) -> str:
    """Fallback PDF parser using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        return _clean_text(text)
    except Exception as e:
        logger.error(f"Fallback PDF parse failed: {e}")
        return ""


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return _clean_text('\n'.join(paragraphs))


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove junk characters."""
    import re
    # Replace multiple whitespace/newlines with single space
    text = re.sub(r'\s+', ' ', text)
    # Remove non-printable chars except common punctuation
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    return text.strip()
